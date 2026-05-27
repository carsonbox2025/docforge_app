#!/usr/bin/env python3
"""生成稿搭子软著说明书 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── 全局样式设置 ───
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页面设置 A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ─── 自定义样式 ───
def set_heading_style(style_name, font_name, font_size, bold=True, color=None, space_before=12, space_after=6):
    s = doc.styles[style_name]
    s.font.name = font_name
    s.font.size = Pt(font_size)
    s.font.bold = bold
    s.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)
    s.paragraph_format.line_spacing = 1.5

set_heading_style('Heading 1', '黑体', 16, space_before=24, space_after=12)
set_heading_style('Heading 2', '黑体', 14, space_before=18, space_after=8)
set_heading_style('Heading 3', '黑体', 13, space_before=12, space_after=6)


def add_para(text, style='Normal', bold=False, font_size=None, alignment=None, font_name=None, color=None, space_before=None, space_after=None, first_line_indent=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font_name or '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or '宋体')
    if bold:
        run.font.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Gray background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing
    return table


def add_screenshot_placeholder(caption_text):
    """添加截图占位框 + 图注"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    # Gray placeholder box
    run = p.add_run('\n\n\n    【此处插入截图】\n\n\n')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Add border around the paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # Gray background
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)

    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption_text)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_body(text):
    """正文段落，首行缩进两字符"""
    return add_para(text, first_line_indent=24, font_size=12)


# ═══════════════════════════════════════════════════════════
# 封面页
# ═══════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

add_para('稿搭子', bold=True, font_size=28, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=4)
add_para('软件操作说明书', bold=True, font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=24)

for _ in range(2):
    doc.add_paragraph()

cover_info = [
    ('软件名称', '稿搭子'),
    ('软件版本', 'V1.0'),
    ('著作权人', '黄卫平'),
    ('开发完成日期', '2026年5月20日'),
    ('首次发表日期', '2026年5月20日'),
    ('编制日期', '2026年5月20日'),
]

for label, value in cover_info:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(f'{label}：')
    run_label.font.name = '宋体'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_label.font.size = Pt(14)
    run_label.font.bold = True
    run_value = p.add_run(value)
    run_value.font.name = '宋体'
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_value.font.size = Pt(14)

# ═══════════════════════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════════════════════
doc.add_page_break()

add_para('目  录', bold=True, font_size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=18)

toc_items = [
    ('1', '引言', 0),
    ('1.1', '编写目的', 1),
    ('1.2', '软件定位', 1),
    ('2', '软件概述', 0),
    ('2.1', '软件简介', 1),
    ('2.2', '功能概览', 1),
    ('2.3', '技术架构', 1),
    ('3', '运行环境', 0),
    ('3.1', '硬件环境', 1),
    ('3.2', '软件环境', 1),
    ('4', '安装与卸载', 0),
    ('4.1', '安装', 1),
    ('4.2', '卸载', 1),
    ('5', '软件功能说明', 0),
    ('5.1', '用户注册与登录', 1),
    ('5.2', '智能生成', 1),
    ('5.3', '文档精修', 1),
    ('5.4', '智能翻译', 1),
    ('5.5', '模板市场', 1),
    ('5.6', '术语管理', 1),
    ('5.7', '文档中心', 1),
    ('5.8', '个人中心', 1),
    ('6', '技术特点', 0),
    ('7', '运行界面', 0),
]

for num, title, level in toc_items:
    indent = '    ' * level
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if level == 0:
        run = p.add_run(f'{indent}{num}  {title}')
        run.font.bold = True
        run.font.size = Pt(12)
    else:
        run = p.add_run(f'{indent}{num}  {title}')
        run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ═══════════════════════════════════════════════════════════
# 第1章 引言
# ═══════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading('1  引言', level=1)
doc.add_heading('1.1  编写目的', level=2)

add_body('本文档是"稿搭子"软件的操作说明书，旨在详细介绍该软件的功能特点、操作流程及技术架构，为计算机软件著作权登记提供必要的技术说明材料。本说明书将全面阐述软件的设计理念、功能模块、技术实现和操作方法，便于著作权审查机构了解软件的核心价值和独创性。')

doc.add_heading('1.2  软件定位', level=2)

add_body('稿搭子是一款基于人工智能技术的跨平台智能文档创作应用，面向企业办公人员、法律从业者、学术研究人员、投标工程师等专业人士，提供从需求描述到专业文档成品的一站式智能创作服务。该软件集成自然语言处理、大语言模型、流式生成等前沿技术，覆盖合同、标书、公文、论文、简历等50余种文档类型的智能生成与处理能力。')

add_body('与传统的文档编辑软件不同，稿搭子以AI大语言模型为核心引擎，用户只需通过自然语言描述创作需求，系统即可自动完成文档结构规划、内容生成、质量审校等全流程工作，大幅降低专业文档创作的技术门槛和时间成本。')

# ═══════════════════════════════════════════════════════════
# 第2章 软件概述
# ═══════════════════════════════════════════════════════════

doc.add_heading('2  软件概述', level=1)
doc.add_heading('2.1  软件简介', level=2)

add_body('稿搭子（DocForge）是一款AI驱动的智能文档创作平台应用，运行于Android移动操作系统。软件以人工智能大语言模型为核心引擎，为用户提供智能文档生成、文档精修审阅、多语种专业翻译、模板套用、术语管理等专业文档处理能力。')

add_body('软件采用Flutter跨平台框架开发，前端使用flutter_riverpod进行声明式状态管理，go_router实现声明式路由导航，dio网络库支持SSE流式通信。后端采用FastAPI异步框架，集成大语言模型API实现智能文档处理能力。软件具备完善的会员订阅体系，支持多渠道支付（华为支付、vivo支付、小米支付、支付宝、微信支付、Apple IAP），为用户提供灵活的服务选择。')

doc.add_heading('2.2  功能概览', level=2)

add_body('本软件具备以下核心功能模块：')

functions = [
    ('智能生成', '用户通过自然语言描述需求，AI自动规划文档结构并流式生成完整文档内容，支持50余种文档类型，覆盖合同、标书、公文、论文、简历、报告等常见场景。'),
    ('文档精修', '对已有文档进行AI智能审阅，自动识别语法错误、优化措辞表达、改善文档结构，以逐条建议卡片的形式呈现审阅结果，支持采纳或拒绝操作。'),
    ('智能翻译', '支持中、英、日、韩、法、德等多语种互译，结合术语表确保专业术语翻译的一致性，完整保留原文排版格式，支持双语对照查看。'),
    ('模板市场', '提供1000余种专业文档模板，覆盖商务、法律、学术、工程等领域，支持分类浏览、关键词搜索、在线预览与一键套用。'),
    ('术语管理', '支持自定义术语条目的增删改查，翻译过程中自动匹配术语表，确保专业用词统一规范，翻译后自动提取新术语。'),
    ('文档中心', '集中管理所有生成和处理的文档，支持按状态分类查看、进度实时追踪、在线内容预览和多格式导出。'),
    ('场景化创作', '针对合同、标书、公文、简历、论文等20余种创作场景提供智能引导和结构化表单输入，降低创作门槛。'),
    ('会员与支付', '提供灵活的会员订阅体系，支持多渠道支付和配额管理，保障服务的可持续运营。'),
]

for i, (name, desc) in enumerate(functions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(f'（{i}）{name}：')
    run_num.font.bold = True
    run_num.font.name = '宋体'
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_num.font.size = Pt(12)
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)

doc.add_heading('2.3  技术架构', level=2)

add_body('本软件采用Flutter跨平台框架开发，基于领域驱动设计（DDD）分层架构组织代码结构。整体技术架构分为以下层次：')

add_table(
    ['层级', '技术选型', '说明'],
    [
        ['UI框架', 'Flutter 3.11+', '跨平台UI框架，一套代码多端运行'],
        ['状态管理', 'flutter_riverpod', '声明式、编译安全的响应式状态管理'],
        ['路由导航', 'go_router', '声明式路由，支持嵌套导航和认证守卫'],
        ['网络通信', 'dio + SSE', 'HTTP客户端，支持SSE流式响应实时推送'],
        ['数据序列化', 'freezed + json_serializable', '不可变数据模型与JSON自动序列化'],
        ['本地存储', 'flutter_secure_storage + hive', '安全凭据存储与轻量键值缓存'],
        ['后端服务', 'FastAPI (Python)', '高性能异步API服务'],
        ['AI引擎', '大语言模型', '文档生成、审阅、翻译的智能处理核心'],
        ['监控系统', 'Sentry + 友盟', '崩溃上报与运营数据分析'],
        ['支付集成', '多渠道SDK', '华为、vivo、小米、支付宝、微信、IAP'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 第3章 运行环境
# ═══════════════════════════════════════════════════════════

doc.add_heading('3  运行环境', level=1)
doc.add_heading('3.1  硬件环境', level=2)

add_table(
    ['项目', '要求'],
    [
        ['处理器', 'ARMv7及以上或ARM64架构处理器'],
        ['内存', '2GB RAM及以上'],
        ['存储空间', '100MB及以上可用存储空间'],
        ['网络', '需要联网使用（Wi-Fi或移动数据）'],
        ['屏幕', '支持4.7英寸及以上屏幕分辨率'],
    ]
)

doc.add_heading('3.2  软件环境', level=2)

add_table(
    ['项目', '要求'],
    [
        ['操作系统', 'Android 8.0（API Level 26）及以上'],
        ['开发框架', 'Flutter 3.11+'],
        ['编程语言', 'Dart 3.11+'],
        ['后端服务', 'FastAPI (Python 3.10+)'],
        ['网络协议', 'HTTPS / SSE'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 第4章 安装与卸载
# ═══════════════════════════════════════════════════════════

doc.add_heading('4  安装与卸载', level=1)
doc.add_heading('4.1  安装', level=2)

add_body('用户可通过以下方式安装稿搭子：')

steps = [
    '应用商店安装：在华为应用市场、小米应用商店、vivo应用商店、OPPO软件商店等主流应用市场搜索"稿搭子"，点击安装即可完成下载和安装。',
    '官网下载安装：访问稿搭子官方网站，下载Android安装包（APK文件），根据系统提示允许安装未知来源应用后完成安装。',
]
for i, step in enumerate(steps, 1):
    add_body(f'（{i}）{step}')

doc.add_heading('4.2  卸载', level=2)

add_body('在设备的"设置→应用管理"中找到"稿搭子"，点击"卸载"即可完成软件卸载。卸载将清除本应用的所有本地缓存数据，但不影响已导出到设备存储的文档文件。')

# ═══════════════════════════════════════════════════════════
# 第5章 软件功能说明
# ═══════════════════════════════════════════════════════════

doc.add_heading('5  软件功能说明', level=1)

# ─── 5.1 用户注册与登录 ───
doc.add_heading('5.1  用户注册与登录', level=2)
doc.add_heading('5.1.1  功能描述', level=3)

add_body('稿搭子提供完整的用户账户体系，支持手机号注册、密码登录、忘记密码重置等功能。新用户首次登录后将进入引导页面，通过分步展示了解软件的智能生成、文档精修、智能翻译等核心功能。系统支持用户资料补全，包括头像设置、昵称修改等个人信息管理。')

add_body('账户体系采用JWT令牌认证机制，用户登录后系统自动维护会话状态，支持令牌自动刷新，确保用户在使用过程中无需重复登录。用户凭证通过flutter_secure_storage安全存储在设备本地，保障账户信息安全。')

doc.add_heading('5.1.2  操作流程', level=3)

login_steps = [
    '打开稿搭子应用，进入启动页面后自动跳转至登录页面。',
    '新用户点击"注册"按钮，输入手机号码并获取短信验证码，设置登录密码完成注册。',
    '已注册用户输入手机号和密码，点击"登录"按钮进入应用主界面。',
    '首次登录后系统展示产品引导页，通过图文介绍智能生成、文档精修、智能翻译等核心功能的使用方法。',
    '用户可在"个人中心"页面补充和完善个人资料信息，包括头像和昵称。',
]
for i, step in enumerate(login_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.1.3  界面示意', level=3)

add_screenshot_placeholder('图5-1 用户登录界面')
add_screenshot_placeholder('图5-2 用户注册界面')
add_screenshot_placeholder('图5-3 新用户引导页')

# ─── 5.2 智能生成 ───
doc.add_heading('5.2  智能生成', level=2)
doc.add_heading('5.2.1  功能描述', level=3)

add_body('智能生成是稿搭子的核心功能模块。用户通过自然语言描述文档创作需求，系统利用AI大语言模型自动分析需求内容、智能规划文档结构、流式生成完整文档内容。整个生成流程采用多阶段流水线架构，分为以下四个阶段：')

gen_stages = [
    ('需求输入阶段', '用户选择文档场景类型（合同、标书、公文、简历、论文等），通过自然语言或结构化表单描述创作需求。支持设置文档语言（中文、英文、日文、韩文）、封面信息、生成模式等参数。系统提供20余种预设场景配置，每个场景包含针对性的表单字段和智能引导。'),
    ('AI规划阶段', '系统自动分析用户的创作需求，基于大语言模型智能规划文档结构大纲，包括章节划分、内容要点分配等。用户可预览大纲结构并进行确认或调整。'),
    ('内容生成阶段', 'AI按照确认的大纲逐章节流式生成文档内容，通过SSE协议实时推送生成内容到客户端。用户可实时查看生成进度，包括当前章节、已完成章节和整体进度百分比。'),
    ('审校微调阶段', '生成完成后用户可在线预览全文内容，系统自动进行质量审校并提示审校发现项（错误、警告、信息三个级别），用户可手动调整内容后选择DOCX、PDF或HTML格式导出文档。'),
]

for name, desc in gen_stages:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run_name = p.add_run(f'● {name}：')
    run_name.font.bold = True
    run_name.font.name = '宋体'
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_name.font.size = Pt(12)
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)

add_body('系统支持两种生成层级：标准文档（Layer 1）适用于常规篇幅的文档创作；长文档（Layer 2）适用于标书、论文等大篇幅文档，支持封面字段动态配置和章节大纲实时追踪，满足不同场景的文档创作需求。')

doc.add_heading('5.2.2  操作流程', level=3)

gen_steps = [
    '在首页点击"智能生成"快捷入口，或通过底部导航栏切换至生成页面。',
    '在需求输入界面，选择文档场景类型（如"合同"、"标书"、"公文"等），系统自动加载对应的表单配置。',
    '根据所选场景填写结构化表单字段（如合同类型、甲方乙方信息等），或直接在文本输入框中输入创作需求描述。',
    '可选设置文档语言、生成模式等高级参数。',
    '点击"开始生成"按钮，系统进入AI规划阶段，自动分析需求并生成文档章节大纲。',
    '用户确认大纲结构后，系统进入内容生成阶段，流式输出各章节内容并实时更新进度。',
    '生成完成后自动进入审校页面，用户可在线预览全文内容，查看系统自动检测的审校发现项。',
    '确认文档内容无误后，点击"导出"按钮，选择DOCX、PDF或HTML格式下载文档到本地。',
]
for i, step in enumerate(gen_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.2.3  界面示意', level=3)

add_screenshot_placeholder('图5-4 智能生成需求输入界面，包含场景选择、表单填写和需求描述')
add_screenshot_placeholder('图5-5 AI规划阶段，展示系统自动生成的文档章节大纲')
add_screenshot_placeholder('图5-6 内容生成阶段，流式输出文档内容并展示章节进度')
add_screenshot_placeholder('图5-7 审校微调页面，展示文档预览和审校发现项')

# ─── 5.3 文档精修 ───
doc.add_heading('5.3  文档精修', level=2)
doc.add_heading('5.3.1  功能描述', level=3)

add_body('文档精修模块为用户提供AI驱动的文档审阅与优化服务。用户可粘贴文本内容或上传已有文档（支持DOCX、TXT、MD格式），系统自动进行多维度审阅分析，以结构化建议卡片的形式逐条展示修改意见。')

add_body('审阅维度涵盖以下五个方面：')

review_dims = [
    ('语法检查', '识别语法错误、用词不当、标点符号问题等基础性错误。'),
    ('风格优化', '改善行文措辞，提升表达流畅度和专业度，优化句式结构。'),
    ('术语规范', '检查专业术语使用的准确性和一致性，确保术语表达规范统一。'),
    ('逻辑连贯', '评估段落间逻辑衔接和论证完整性，发现论证缺失或逻辑跳跃问题。'),
    ('格式规范', '检查标题层级、编号格式、排版规范等格式层面的问题。'),
]

for name, desc in review_dims:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    run_name = p.add_run(f'● {name}：')
    run_name.font.bold = True
    run_name.font.name = '宋体'
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_name.font.size = Pt(12)
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)

add_body('每条审阅建议包含：原文内容、修改建议、修改原因说明和严重程度标识（错误/警告/建议）四个要素。用户可逐条查看建议并选择"采纳"或"拒绝"，也支持"全部采纳"和"全部拒绝"的批量操作。审阅结果页面支持宽屏双栏布局（左侧文档预览、右侧建议面板）和窄屏单栏布局自适应切换。')

add_body('系统支持三种润色强度：轻度润色仅修正明确错误，保留原文风格；中度润色修正错误并优化表达，提升文档质量；深度润色全面优化文档结构、措辞和逻辑连贯性。同时支持文本精修和文档精修两种输入模式，文档精修模式下系统自动识别文档类型并针对性优化。')

add_body('审阅完成后，用户可选择"确认导出"生成精修版文档，或选择"审阅报告"导出包含所有修改记录的审阅报告，两种模式均支持DOCX、PDF、HTML格式导出。')

doc.add_heading('5.3.2  操作流程', level=3)

polish_steps = [
    '在底部导航栏或首页点击"文档精修"快捷入口。',
    '选择输入模式："文本精修"（粘贴文本内容）或"文档精修"（上传文件）。',
    '输入或上传需要审阅的文档内容。文档精修模式支持DOCX、TXT、MD格式文件上传。',
    '选择文档类型（合同、公文、论文、标书、简历、报告、会议纪要、技术方案、诉讼文书、通用），帮助AI更精准地审阅。',
    '选择润色强度：轻度、中度或深度。',
    '点击"开始审阅"按钮，系统进入AI分析阶段，实时展示审阅进度和已发现的建议条目。',
    '审阅完成后进入审阅结果页面，系统以建议卡片列表形式展示所有修改建议。',
    '点击建议卡片可定位到文档对应段落，在预览区域查看上下文。',
    '逐条处理建议：点击"采纳"应用修改，点击"拒绝"保留原文。也可使用"全部采纳"或"全部拒绝"批量处理。',
    '所有建议处理完毕后，点击底部"确认导出"生成精修文档，或点击"审阅报告"导出审阅记录。',
]
for i, step in enumerate(polish_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.3.3  界面示意', level=3)

add_screenshot_placeholder('图5-8 文档精修输入界面，展示输入模式切换、文件上传区域和参数配置')
add_screenshot_placeholder('图5-9 AI审阅进行中，展示章节审阅进度和实时发现的建议')
add_screenshot_placeholder('图5-10 审阅结果页面，展示建议卡片列表和统计信息')
add_screenshot_placeholder('图5-11 宽屏双栏布局，左侧文档预览与右侧建议面板联动')

# ─── 5.4 智能翻译 ───
doc.add_heading('5.4  智能翻译', level=2)
doc.add_heading('5.4.1  功能描述', level=3)

add_body('智能翻译模块提供专业的多语种文档翻译服务，支持中文、英语、日语、韩语、法语、德语等多种语言之间的互译。该模块具备以下核心能力：')

trans_features = [
    ('双模式翻译', '支持文本翻译（粘贴文本内容即时翻译）和文档翻译（上传完整文档批量翻译）两种工作模式，满足不同翻译场景需求。'),
    ('术语一致性保障', '翻译过程中自动匹配用户术语表中的专业术语条目，确保术语翻译在整篇文档中的准确性和一致性。翻译完成后展示术语匹配统计信息，包括匹配术语数量和一致性校验结果。'),
    ('多维度翻译配置', '支持选择文档类型（合同、标书、论文、报告等）和行业领域，AI根据文档上下文进行精准翻译。支持用户输入自定义翻译要求（如"保持专业术语不翻译"、"使用正式语气"等），实现个性化翻译。'),
    ('双语对照查看', '翻译结果支持"原文"、"译文"、"双语"三种查看模式。双语模式以段落为单位逐段对照展示原文与译文，便于翻译校对。译文中自动高亮匹配术语表的专业术语，方便用户快速识别。'),
    ('格式保留与导出', '文档翻译模式完整保留原文的排版格式和文档结构，包括标题层级、列表格式、段落间距等。翻译结果支持导出为DOCX、PDF、HTML三种格式。'),
]

for name, desc in trans_features:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run_name = p.add_run(f'● {name}：')
    run_name.font.bold = True
    run_name.font.name = '宋体'
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_name.font.size = Pt(12)
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)

doc.add_heading('5.4.2  操作流程', level=3)

trans_steps = [
    '在底部导航栏或首页点击"智能翻译"快捷入口。',
    '选择翻译模式："文本翻译"或"文档翻译"。',
    '选择源语言和目标语言。点击语言标签可打开语言选择面板，点击中间的语言互换按钮可快速交换源语言和目标语言。',
    '在文本模式下输入或粘贴需要翻译的文本内容；在文档模式下上传DOCX、PDF、TXT、MD格式的文档文件。',
    '选择文档类型（合同、公文、论文、报告等）和行业领域（可选），帮助AI进行更精准的上下文翻译。',
    '输入自定义翻译要求（可选），如特定术语处理规则、语气风格要求等。',
    '查看术语表中的术语条目，可点击"编辑"跳转至术语管理页面维护术语库。',
    '点击"专业翻译"按钮开始翻译。翻译过程中可查看实时进度、段落翻译状态和术语提取信息。',
    '翻译完成后，在结果页面切换"译文"、"原文"、"双语"三种查看模式查看翻译结果。',
    '确认翻译质量后，点击"导出文档"按钮选择DOCX、PDF或HTML格式下载翻译结果。',
]
for i, step in enumerate(trans_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.4.3  界面示意', level=3)

add_screenshot_placeholder('图5-12 文本翻译模式，展示语言选择器、文本输入区域和术语表')
add_screenshot_placeholder('图5-13 文档翻译模式，展示文件上传区域和翻译参数配置')
add_screenshot_placeholder('图5-14 翻译进度界面，展示段落翻译进度和术语提取统计')
add_screenshot_placeholder('图5-15 翻译结果页面，展示术语统计、查看模式切换和导出按钮')
add_screenshot_placeholder('图5-16 双语对照查看模式，逐段对照展示原文与译文')

# ─── 5.5 模板市场 ───
doc.add_heading('5.5  模板市场', level=2)
doc.add_heading('5.5.1  功能描述', level=3)

add_body('模板市场模块为用户提供丰富的专业文档模板库，包含1000余种精选模板，覆盖商务、法律、学术、工程等多个领域。模板按类别分类组织，用户可浏览、搜索、预览并一键套用模板生成标准化文档。')

add_body('主要功能包括：分类浏览（模板按商务、法律、学术、工程等类别分类，通过顶部标签栏快速切换筛选）、关键词搜索（支持输入关键词快速定位目标模板）、模板预览（点击模板卡片进入预览页面查看模板结构和示例效果）、一键套用（选择模板后直接进入智能生成流程，AI基于模板结构和用户输入生成文档）、Pro专属模板（高级会员可解锁专业级模板库，未订阅用户点击时展示会员升级引导）。')

doc.add_heading('5.5.2  操作流程', level=3)

tpl_steps = [
    '在首页点击"模板市场"快捷入口进入模板市场页面。',
    '通过顶部分类标签栏浏览不同类别的模板，或使用搜索功能输入关键词查找模板。',
    '点击模板卡片查看模板详情和在线预览效果。',
    '点击"使用模板"按钮，系统自动进入智能生成流程，加载该模板对应的场景配置。',
    '填写场景表单信息后，AI基于模板结构生成完整文档。',
]
for i, step in enumerate(tpl_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.5.3  界面示意', level=3)

add_screenshot_placeholder('图5-17 模板市场主界面，展示分类标签和模板卡片网格布局')
add_screenshot_placeholder('图5-18 模板预览页面，展示模板详情和示例内容')

# ─── 5.6 术语管理 ───
doc.add_heading('5.6  术语管理', level=2)
doc.add_heading('5.6.1  功能描述', level=3)

add_body('术语管理模块提供专业术语库的管理工具，用户可创建和维护自定义术语表，确保文档翻译和专业写作中的用词统一规范。该模块支持术语条目的新增、编辑和删除操作，每条术语包含源语言术语、目标语言术语、所属语言对和备注信息。')

add_body('系统支持中文-英文、中文-日文、中文-韩文、英文-中文、英文-日文等多种语言对配置。在智能翻译过程中，系统自动匹配术语表中的条目，确保专业术语翻译的一致性。翻译完成后系统自动提取新发现的术语，并提示用户保存到术语库，持续丰富术语资源。')

doc.add_heading('5.6.2  操作流程', level=3)

glossary_steps = [
    '通过翻译页面的"编辑术语表"链接或个人中心进入术语管理页面。',
    '点击"新增术语"按钮，在弹出的表单中输入源术语、目标术语，选择所属语言对，填写备注说明。',
    '保存后术语条目出现在术语列表中，可随时编辑或删除。',
    '在智能翻译过程中，系统自动引用术语表进行翻译，确保术语一致性。',
    '翻译完成后，系统提示自动提取的新术语，用户可确认保存到术语库。',
]
for i, step in enumerate(glossary_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.6.3  界面示意', level=3)

add_screenshot_placeholder('图5-19 术语管理页面，展示术语条目列表和语言对标识')

# ─── 5.7 文档中心 ───
doc.add_heading('5.7  文档中心', level=2)
doc.add_heading('5.7.1  功能描述', level=3)

add_body('文档中心模块是用户管理所有文档的工作台，集中展示通过智能生成、文档精修、智能翻译等功能创建的所有文档记录。模块支持按文档处理状态分类查看、生成进度实时追踪、文档内容在线预览和多格式导出等核心功能。')

add_body('文档按处理状态分为三类：进行中（正在生成或处理的文档，实时展示进度百分比和进度条）、已完成（处理完成的文档，可预览、导出和进一步编辑）、失败（处理过程中出现错误的文档，展示错误信息并支持重试操作）。')

add_body('文档列表支持下拉刷新获取最新状态、滚动到底部自动加载更多文档的分页加载机制。通过顶部标签栏可按"进行中"、"已完成"、"全部"三种维度筛选查看。点击文档卡片可进入文档详情页面，查看文档完整内容并进行导出操作。')

doc.add_heading('5.7.2  操作流程', level=3)

doc_steps = [
    '点击底部导航栏"文档"入口进入文档中心页面。',
    '通过顶部标签栏切换查看"进行中"、"已完成"、"全部"分类的文档列表。',
    '下拉刷新获取最新文档状态，向上滚动到底部自动加载更多历史文档。',
    '点击文档卡片进入文档详情页面，查看文档完整内容。',
    '在文档详情页面，选择DOCX、PDF或HTML格式导出文档到本地。',
]
for i, step in enumerate(doc_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.7.3  界面示意', level=3)

add_screenshot_placeholder('图5-20 文档中心主界面，展示按状态分类的文档列表和进度展示')
add_screenshot_placeholder('图5-21 文档详情页面，展示文档内容预览和导出操作')

# ─── 5.8 个人中心 ───
doc.add_heading('5.8  个人中心', level=2)
doc.add_heading('5.8.1  功能描述', level=3)

add_body('个人中心模块提供用户信息管理、会员订阅管理、使用统计查询、消息通知查看、应用设置等功能。用户可在此查看和编辑个人头像、昵称等基本资料信息，查看当前会员状态和权益信息，订阅Pro会员解锁更多高级功能和模板资源。')

add_body('使用统计功能展示用户的文档生成次数、精修次数、翻译次数等用量数据，帮助用户了解自身使用情况。消息通知模块接收系统通知和重要消息提醒。用户反馈功能支持用户提交产品使用反馈和问题报告，帮助持续改进产品体验。应用设置包含关于稿搭子版本信息、隐私政策、用户协议等法律合规文档。')

doc.add_heading('5.8.2  操作流程', level=3)

profile_steps = [
    '点击底部导航栏"我的"入口进入个人中心页面。',
    '查看个人基本信息和当前会员状态。',
    '点击相应功能菜单项进入对应功能页面（会员管理、使用统计、消息通知等）。',
    '在"关于"页面查看应用版本信息和相关法律文档。',
]
for i, step in enumerate(profile_steps, 1):
    add_body(f'{i}. {step}')

doc.add_heading('5.8.3  界面示意', level=3)

add_screenshot_placeholder('图5-22 个人中心主界面，展示用户信息、会员状态和功能菜单')

# ═══════════════════════════════════════════════════════════
# 第6章 技术特点
# ═══════════════════════════════════════════════════════════

doc.add_heading('6  技术特点', level=1)

doc.add_heading('6.1  跨平台架构', level=2)
add_body('稿搭子采用Flutter跨平台框架开发，实现Android、iOS、Web、Windows、macOS多端统一代码库，确保各平台用户体验一致性。核心业务逻辑与UI表现层完全分离，通过flutter_riverpod进行声明式状态管理，确保数据流向清晰可控，提高代码的可维护性和可测试性。')

doc.add_heading('6.2  AI智能文档处理引擎', level=2)
add_body('软件集成大语言模型能力，采用SSE（Server-Sent Events）流式通信协议，实现文档内容的实时流式生成和进度追踪。AI处理流程包含需求分析、结构规划、内容生成、质量审校等多个智能阶段，每个阶段均可独立优化和扩展，确保输出文档的专业性和完整性。')

doc.add_heading('6.3  领域驱动分层架构', level=2)
add_body('软件采用领域驱动设计（DDD）分层架构组织代码结构，将每个功能模块划分为data（数据层）、domain（领域层）、presentation（表现层）三个层次。数据层负责网络请求封装和本地数据持久化，领域层封装业务逻辑和状态管理，表现层专注UI组件构建和用户交互处理。这种分层架构确保了代码的高内聚低耦合，提升了系统的可维护性和可扩展性。')

doc.add_heading('6.4  多渠道支付集成', level=2)
add_body('软件实现了统一的多渠道支付系统，根据运行环境自动选择最优支付通道。在Android端集成华为支付SDK、vivo支付SDK、小米支付SDK，Web端支持支付宝和微信支付，iOS端支持Apple IAP内购。支付流程包含商品查询、下单、支付确认、发货验证、掉单恢复的完整闭环链路，确保交易安全可靠。')

doc.add_heading('6.5  术语一致性引擎', level=2)
add_body('软件内置专业术语管理引擎，支持用户自定义术语表的创建和维护。在翻译和文档生成过程中，系统自动匹配术语表中的条目进行替换和校验，确保专业术语在整个文档中的翻译一致性和用词规范性。翻译完成后系统自动提取新发现的术语并提示用户保存，实现术语库的持续积累和优化。')

doc.add_heading('6.6  数据安全保障', level=2)
add_body('软件采用flutter_secure_storage进行敏感凭据（如认证令牌）的安全存储，使用Hive轻量级数据库进行非敏感数据的本地缓存。所有网络通信均采用HTTPS加密传输，用户认证基于JWT令牌机制并支持自动刷新，确保用户数据的传输安全和存储安全。')

# ═══════════════════════════════════════════════════════════
# 第7章 运行界面
# ═══════════════════════════════════════════════════════════

doc.add_heading('7  运行界面', level=1)

add_body('以下展示稿搭子软件在Android设备上的主要运行界面截图，涵盖首页、智能生成、文档精修、智能翻译、模板市场、文档中心和个人中心等核心功能模块。')

add_screenshot_placeholder('图7-1 应用启动画面')
add_screenshot_placeholder('图7-2 首页界面，展示轮播图、快捷操作入口和最近文档')
add_screenshot_placeholder('图7-3 智能生成完整流程示意')
add_screenshot_placeholder('图7-4 文档精修完整流程示意')
add_screenshot_placeholder('图7-5 智能翻译完整流程示意')
add_screenshot_placeholder('图7-6 模板市场浏览界面')
add_screenshot_placeholder('图7-7 文档中心管理界面')
add_screenshot_placeholder('图7-8 个人中心与设置界面')

# ─── 结尾信息 ───
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 稿搭子 V1.0 软件操作说明书 · 完 —')
run.font.size = Pt(12)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('著作权人：黄卫平')
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('编制完成日期：2026年5月20日')
run.font.size = Pt(11)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ─── 保存 ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '稿搭子_软件操作说明书_V1.0.docx')
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
